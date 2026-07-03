"""Document translation worker.

Consumes jobs from the Redis `doc-jobs` queue (payload: docs/04-API.md §8),
translates documents with Claude, stores results in MinIO under
jobs/<jobId>/result/, and reports status back to the API via
PUT /api/documents/<jobId>/status (authenticated with WORKER_TOKEN).

Supported inputs:
- DOCX  — paragraph-by-paragraph translation, formatting preserved (Phase 2 core).
- PDF   — text extraction via PyMuPDF, translated into a clean DOCX.
- JPG/PNG — Tesseract OCR, translated into a clean DOCX.
"""

import asyncio
import io
import json
import os
import sys
from urllib.parse import urlparse

import aiohttp
import docx
from anthropic import AsyncAnthropic
from minio import Minio
from redis.asyncio import Redis

REDIS_URL = os.getenv("REDIS_URL", "redis://localhost:6379/0")
API_BASE_URL = os.getenv("API_BASE_URL", "http://localhost:3000/api")
WORKER_TOKEN = os.getenv("WORKER_TOKEN")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
ANTHROPIC_MODEL = os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6")

S3_ENDPOINT = os.getenv("S3_ENDPOINT", "http://localhost:9000")
S3_ACCESS_KEY = os.getenv("S3_ACCESS_KEY", "")
S3_SECRET_KEY = os.getenv("S3_SECRET_KEY", "")
S3_BUCKET = os.getenv("S3_BUCKET", "lingo-files")

QUEUE = "doc-jobs"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

if not WORKER_TOKEN or not ANTHROPIC_API_KEY:
    print("WORKER_TOKEN and ANTHROPIC_API_KEY are required", file=sys.stderr)
    sys.exit(1)

_endpoint = urlparse(S3_ENDPOINT)
ai_client = AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
redis_client = Redis.from_url(REDIS_URL)
s3_client = Minio(
    _endpoint.netloc,
    access_key=S3_ACCESS_KEY,
    secret_key=S3_SECRET_KEY,
    secure=_endpoint.scheme == "https",
)

LANG_NAMES = {"UZ": "Uzbek (Latin script)", "EN": "English", "RU": "Russian"}


async def translate_text(text: str, from_lang: str, to_lang: str) -> str:
    if not text.strip():
        return text
    resp = await ai_client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=4096,
        system=(
            "You are a certified document translator. Translate exactly, preserving "
            "numbers, names, dates and formatting markers. Output only the translation."
        ),
        messages=[
            {
                "role": "user",
                "content": (
                    f"Translate from {LANG_NAMES.get(from_lang, from_lang)} "
                    f"to {LANG_NAMES.get(to_lang, to_lang)}:\n\n{text}"
                ),
            }
        ],
    )
    return resp.content[0].text


async def translate_long_text(text: str, from_lang: str, to_lang: str, chunk_chars: int = 6000) -> str:
    """Translate arbitrary-length text by paragraph-aligned chunks."""
    paragraphs = text.split("\n\n")
    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) > chunk_chars and current:
            chunks.append(current)
            current = para
        else:
            current = f"{current}\n\n{para}" if current else para
    if current:
        chunks.append(current)

    translated = [await translate_text(chunk, from_lang, to_lang) for chunk in chunks]
    return "\n\n".join(translated)


def fetch_object(key: str) -> bytes:
    resp = s3_client.get_object(S3_BUCKET, key)
    try:
        return resp.read()
    finally:
        resp.close()
        resp.release_conn()


def store_object(key: str, data: bytes, mime_type: str) -> None:
    s3_client.put_object(S3_BUCKET, key, io.BytesIO(data), len(data), content_type=mime_type)


def text_to_docx(text: str) -> bytes:
    document = docx.Document()
    for para in text.split("\n"):
        document.add_paragraph(para)
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


async def translate_docx(file_bytes: bytes, from_lang: str, to_lang: str) -> bytes:
    document = docx.Document(io.BytesIO(file_bytes))
    for para in document.paragraphs:
        if para.text.strip():
            para.text = await translate_text(para.text, from_lang, to_lang)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = await translate_text(cell.text, from_lang, to_lang)
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def extract_pdf_text(file_bytes: bytes) -> str:
    import fitz  # PyMuPDF

    with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
        return "\n\n".join(page.get_text() for page in pdf)


def ocr_image(file_bytes: bytes) -> str:
    import pytesseract
    from PIL import Image

    langs = os.getenv("TESSERACT_LANGS", "eng+rus+uzb")
    return pytesseract.image_to_string(Image.open(io.BytesIO(file_bytes)), lang=langs)


async def process_source_file(src: dict, job: dict) -> dict:
    """Translate one source file; returns a resultFiles entry for the API callback."""
    file_bytes = fetch_object(src["storageKey"])
    from_lang, to_lang = job["fromLang"], job["toLang"]
    mime = src["mimeType"]

    if mime == DOCX_MIME:
        result_bytes = await translate_docx(file_bytes, from_lang, to_lang)
    elif mime == "application/pdf":
        text = extract_pdf_text(file_bytes)
        result_bytes = text_to_docx(await translate_long_text(text, from_lang, to_lang))
    elif mime in ("image/jpeg", "image/png"):
        text = ocr_image(file_bytes)
        result_bytes = text_to_docx(await translate_long_text(text, from_lang, to_lang))
    else:
        raise ValueError(f"Unsupported mime type: {mime}")

    original_stem = os.path.splitext(os.path.basename(src["storageKey"]))[0]
    result_name = f"{original_stem}_{to_lang.lower()}.docx"
    result_key = f"jobs/{job['jobId']}/result/{result_name}"
    store_object(result_key, result_bytes, DOCX_MIME)

    return {
        "storageKey": result_key,
        "originalName": result_name,
        "mimeType": DOCX_MIME,
        "sizeBytes": len(result_bytes),
    }


async def notify_api(session: aiohttp.ClientSession, job_id: str, payload: dict) -> None:
    async with session.put(
        f"{API_BASE_URL}/documents/{job_id}/status",
        json=payload,
        headers={"x-worker-token": WORKER_TOKEN},
    ) as resp:
        if resp.status >= 400:
            body = await resp.text()
            print(f"API status callback failed ({resp.status}): {body}", file=sys.stderr)


async def process_job(job: dict) -> None:
    job_id = job["jobId"]
    async with aiohttp.ClientSession() as session:
        try:
            await notify_api(session, job_id, {"status": "PROCESSING"})

            result_files = [await process_source_file(src, job) for src in job["sourceFiles"]]

            final_status = "REVIEW" if job.get("notarize") else "DONE"
            await notify_api(
                session, job_id, {"status": final_status, "resultFiles": result_files}
            )
            print(f"Job {job_id} finished: {final_status}, {len(result_files)} file(s)")
        except Exception as exc:  # report failure, never leave a job hanging
            print(f"Job {job_id} failed: {exc}", file=sys.stderr)
            await notify_api(
                session,
                job_id,
                {"status": "FAILED", "errorMessage": str(exc)[:2000]},
            )


async def main() -> None:
    print(f"doc-worker started. queue={QUEUE} model={ANTHROPIC_MODEL} bucket={S3_BUCKET}")
    while True:
        try:
            item = await redis_client.blpop(QUEUE, timeout=5)
            if item:
                await process_job(json.loads(item[1]))
        except Exception as exc:
            print(f"Worker loop error: {exc}", file=sys.stderr)
            await asyncio.sleep(2)


if __name__ == "__main__":
    asyncio.run(main())
